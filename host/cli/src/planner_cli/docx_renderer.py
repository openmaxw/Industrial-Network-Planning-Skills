from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from planner_cli.document_model import Appendix, BulletBlock, DiagramBlock, DocumentModel, ParagraphBlock, Section, TableBlock
from planner_cli.graphviz_support import graphviz_install_hint, has_graphviz
from planner_cli.topology_diagrams import render_diagram_png


def _ensure_styles(document: Document) -> None:
    styles = document.styles
    names = [s.name for s in styles]
    if "Diagram Code" not in names:
        style = styles.add_style("Diagram Code", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8)
    if "Cover Meta" not in names:
        style = styles.add_style("Cover Meta", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(12)
    if "Result Paragraph" not in names:
        style = styles.add_style("Result Paragraph", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(11)


def _set_document_defaults(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _add_header_footer(document: Document, model: DocumentModel) -> None:
    for section in document.sections:
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.text = f"{model.meta.project_name}｜正式交付件"
        if p.runs:
            _set_run_font(p.runs[0], size=9, color=(71, 85, 105))

        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = "第 "
        _add_page_number(fp)
        tail = fp.add_run(" 页")
        _set_run_font(tail, size=9, color=(100, 116, 139))


def _add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '目录将在 Word 中打开后自动更新。'
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def _add_title_page(document: Document, model: DocumentModel) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    p.space_before = Pt(100)
    run = p.add_run(model.meta.project_name or model.meta.document_title or "网络规划方案")
    _set_run_font(run, size=22, bold=True, color=(15, 23, 42))

    subtitle = document.add_paragraph(style="Cover Meta")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("工业网络正式方案").bold = True

    for text in [
        f"建设单位：{model.meta.customer_name}" if model.meta.customer_name else "",
        f"项目地点：{model.meta.site_name}" if model.meta.site_name else "",
        f"项目目标：{model.meta.objective}" if model.meta.objective else "",
        f"范围说明：{model.meta.scope}" if model.meta.scope else "",
        f"版本：{model.meta.version}",
    ]:
        if text:
            para = document.add_paragraph(style="Cover Meta")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(text)

    document.add_page_break()
    document.add_heading("目录", level=1)
    toc_para = document.add_paragraph()
    _add_toc(toc_para)
    document.add_page_break()


def _render_paragraph(document: Document, block: ParagraphBlock) -> None:
    para = document.add_paragraph(style="Result Paragraph")
    text = block.text
    if text.startswith("【") and "】" in text:
        prefix, body = text.split("】", 1)
        r1 = para.add_run(prefix + "】")
        _set_run_font(r1, size=11, bold=True, color=(30, 64, 175))
        r2 = para.add_run(body)
        _set_run_font(r2, size=11)
    else:
        r = para.add_run(text)
        _set_run_font(r, size=11)


def _render_bullets(document: Document, block: BulletBlock) -> None:
    if block.title:
        document.add_heading(block.title, level=3)
    for item in block.items:
        document.add_paragraph(item, style="List Bullet")


def _render_table(document: Document, block: TableBlock) -> None:
    if block.title:
        document.add_heading(block.title, level=3)
    table = document.add_table(rows=len(block.rows) + 1, cols=len(block.headers))
    table.style = "Table Grid"
    for index, header in enumerate(block.headers):
        cell = table.cell(0, index)
        cell.text = str(header)
    for row_index, row in enumerate(block.rows, start=1):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = str(value)
    document.add_paragraph("")


def _render_diagram(document: Document, block: DiagramBlock) -> None:
    if block.title:
        document.add_heading(block.title, level=3)
    if block.source_format == 'graphviz' and not has_graphviz():
        document.add_paragraph(graphviz_install_hint())
        for line in block.content.splitlines():
            document.add_paragraph(line, style="Diagram Code")
        return
    with tempfile.TemporaryDirectory(prefix="planner-diagram-") as tmpdir:
        image_path = Path(tmpdir) / f"{block.diagram_type or 'diagram'}.png"
        try:
            render_diagram_png(block, image_path)
            pic = document.add_picture(str(image_path), width=Inches(7.6))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if block.caption:
                cp = document.add_paragraph(block.caption)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        except Exception:
            pass
    document.add_paragraph("拓扑图未能生成图片，以下保留源代码以便继续渲染：")
    for line in block.content.splitlines():
        document.add_paragraph(line, style="Diagram Code")


def _render_section(document: Document, section: Section, level: int = 1) -> None:
    document.add_heading(section.title, level=min(level, 2))
    if section.summary:
        document.add_paragraph(section.summary)
    for block in section.paragraphs:
        _render_paragraph(document, block)
    for block in section.tables:
        _render_table(document, block)
    for block in section.diagrams:
        _render_diagram(document, block)
    for block in section.bullets:
        _render_bullets(document, block)


def _render_appendix(document: Document, appendix: Appendix) -> None:
    document.add_page_break()
    document.add_heading(appendix.title, level=1)
    for block in appendix.paragraphs:
        _render_paragraph(document, block)
    for block in appendix.tables:
        _render_table(document, block)
    for block in appendix.diagrams:
        _render_diagram(document, block)


def render_docx(model: DocumentModel, output_path: Path) -> None:
    document = Document()
    _ensure_styles(document)
    _set_document_defaults(document)
    _add_title_page(document, model)
    for section in model.sections:
        _render_section(document, section)
    for appendix in model.appendices:
        _render_appendix(document, appendix)
    _add_header_footer(document, model)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
